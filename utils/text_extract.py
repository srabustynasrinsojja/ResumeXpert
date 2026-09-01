from __future__ import annotations

from pathlib import Path
from typing import List
import re

import pdfplumber
from docx import Document


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}


def clean_extracted_text(text: str) -> str:
    text = str(text or "")
    text = text.replace("\x00", " ")
    text = re.sub(r"\r", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


def extract_text_from_pdf(file_path: str | Path) -> str:
    file_path = Path(file_path)
    pages: List[str] = []

    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            if text.strip():
                pages.append(text)

    return clean_extracted_text("\n\n".join(pages))


def extract_text_from_docx(file_path: str | Path) -> str:
    file_path = Path(file_path)
    doc = Document(file_path)
    chunks: List[str] = []

    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            chunks.append(t)

    for table in doc.tables:
        for row in table.rows:
            cells = []
            for cell in row.cells:
                t = (cell.text or "").strip()
                if t:
                    cells.append(t)
            if cells:
                chunks.append(" | ".join(cells))

    return clean_extracted_text("\n".join(chunks))


def extract_text_from_txt(file_path: str | Path) -> str:
    file_path = Path(file_path)
    return clean_extracted_text(file_path.read_text(encoding="utf-8", errors="ignore"))


def extract_text_from_file(file_path: str | Path) -> str:
    file_path = Path(file_path)
    ext = file_path.suffix.lower()

    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {ext}. Supported: {sorted(SUPPORTED_EXTENSIONS)}")

    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    if ext == ".docx":
        return extract_text_from_docx(file_path)
    if ext == ".txt":
        return extract_text_from_txt(file_path)

    raise ValueError(f"Unhandled file type: {ext}")
